#!/usr/bin/env python3
"""Generate Lumio AI × Morgan Stanley one-pager as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ---------- page setup ----------
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

NAVY   = RGBColor(0x0A, 0x16, 0x28)
ACCENT = RGBColor(0x1A, 0x36, 0x5D)
GOLD   = RGBColor(0x8B, 0x69, 0x14)
GRAY   = RGBColor(0x55, 0x55, 0x55)

# ---------- helpers ----------
def add_heading_styled(text, level=1, color=ACCENT):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, italic=False, size=Pt(10.5), color=None, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_rich_para(segments, align=None, space_after=Pt(6)):
    """segments = list of (text, bold, italic, color_or_None)"""
    p = doc.add_paragraph()
    for text, bold, italic, color in segments:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        if color:
            run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def set_cell_shading(cell, hex_color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): hex_color,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): '4',
            qn('w:space'): '0',
            qn('w:color'): 'CCCCCC',
        })
        borders.append(element)
    tblPr.append(borders)

# ============================================================
# HEADER
# ============================================================
title = doc.add_heading('Lumio AI — Project Summary', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    run.font.size = Pt(22)

add_para('Bank-Grade AI Execution & Governance Layer for Enterprise Workflows',
         italic=True, size=Pt(11), color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

add_rich_para([
    ('Prepared by: ', True, False, GRAY),
    ('Songyi Li  ·  ucessl7@ucl.ac.uk', False, False, ACCENT),
], align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))

add_para('26 February 2026', size=Pt(9), color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))

# Divider
doc.add_paragraph('─' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# THE IDEA
# ============================================================
add_heading_styled('The Idea', level=1)

add_rich_para([
    ('Lumio AI', True, False, None),
    (' is a bank-grade AI execution and governance platform designed to bridge the gap between ', False, False, None),
    ('AI that can be demonstrated', False, True, None),
    (' and ', False, False, None),
    ('AI that can be deployed in production', False, True, None),
    (' — safely, transparently, and at scale.', False, False, None),
])

add_rich_para([
    ('At its core, Lumio provides three capabilities that large financial institutions need but rarely find in a single solution:', False, False, None),
])

capabilities = [
    ('Super Agent Orchestration', 'Explainable, multi-agent workflow execution with dynamic task routing, failure recovery, and rollback paths.'),
    ('Evidence & Audit Infrastructure', 'Every AI-generated output is linked to its sources, decisions are traceable, and full execution logs are immutable and export-ready for compliance review.'),
    ('Digital Twin Optimisation Engine (DTOE)', 'A continuously-learning "belief state" model of each client or user, powered by Bellman/MPC-style optimisation, that generates personalised next-best-action recommendations and improves over time through outcome feedback loops.'),
]

for i, (cap_title, cap_desc) in enumerate(capabilities, 1):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{i}. ')
    run_num.font.name = 'Calibri'
    run_num.font.size = Pt(10.5)
    run_title = p.add_run(cap_title)
    run_title.bold = True
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(10.5)
    run_desc = p.add_run(f' — {cap_desc}')
    run_desc.font.name = 'Calibri'
    run_desc.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(4)

add_rich_para([
    ('All of this is wrapped in a ', False, False, None),
    ('human-in-the-loop control plane', True, False, None),
    (': high-risk actions require explicit approval, policies are enforced by domain and risk level, and the system can be rolled back to conservative mode at any time.', False, False, None),
], space_after=Pt(12))

# ============================================================
# WHY MORGAN STANLEY
# ============================================================
add_heading_styled('Why Morgan Stanley', level=1)

add_rich_para([
    ('Morgan Stanley is already operating at AI scale — from ', False, False, None),
    ('AI Debrief', True, False, None),
    (' in Wealth Management to ', False, False, None),
    ('AskResearchGPT', True, False, None),
    (' across Institutional Securities. Lumio AI is not positioned to replace these capabilities. Instead, it fills the operational gaps that emerge as AI moves from pilot to production:', False, False, None),
])

# Challenge table
table_data = [
    ('Workflow Reliability', 'End-to-end orchestration with structured task graphs, evidence gates, and automatic fallback.'),
    ('Governance Consistency', 'Policy-driven routing, mandatory human checkpoints for sensitive outputs, and standardised quality gates across business lines.'),
    ('Auditability', 'Immutable execution traces, citation-linked outputs, and one-click audit export — aligned with Morgan Stanley\'s NIST-aligned cybersecurity and third-party governance frameworks.'),
    ('Personalisation at Scale', 'DTOE-powered next-best-action ranking for advisors, dynamically adapting recommendations based on client context, market conditions, and real outcome feedback.'),
]

table = doc.add_table(rows=1, cols=2)
style_table(table)
hdr = table.rows[0].cells
hdr[0].text = 'Challenge'
hdr[1].text = "Lumio's Contribution"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
    set_cell_shading(cell, '1A365D')

for challenge, contribution in table_data:
    row = table.add_row().cells
    row[0].text = challenge
    row[1].text = contribution
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
    row[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()  # spacer

add_heading_styled('Priority Use Cases', level=2)

use_cases = [
    ('Wealth Management', 'Pre-meeting prep → post-meeting notes → CRM-ready draft → compliance check, reducing advisor admin time and improving CRM completeness.'),
    ('Institutional Securities', 'Research retrieval → client-specific brief → citation verification → outbound draft, accelerating analyst throughput and eliminating citation gaps.'),
    ('Risk & Compliance', 'AI use-case intake → evaluation → approval → regression monitoring, shortening approval cycles and increasing validation coverage.'),
    ('Advisory Personalisation', 'Twin-aware next-best-action selection per client context, boosting recommendation adoption and reducing manual rewrite.'),
]

for uc_title, uc_desc in use_cases:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run_t = p.add_run(uc_title)
    run_t.bold = True
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(10)
    run_d = p.add_run(f' — {uc_desc}')
    run_d.font.name = 'Calibri'
    run_d.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

doc.add_paragraph()

# ============================================================
# 90-DAY PILOT
# ============================================================
add_heading_styled('Proposed Engagement: 90-Day Pilot', level=1)

pilot_table = doc.add_table(rows=1, cols=3)
style_table(pilot_table)
hdr = pilot_table.rows[0].cells
hdr[0].text = 'Phase'
hdr[1].text = 'Timeline'
hdr[2].text = 'Activities'
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
    set_cell_shading(cell, '1A365D')

pilot_data = [
    ('Scope & Controls', 'Weeks 1–2', 'Select one P1 workflow; define data boundaries, approval checkpoints, audit fields, and success metrics.'),
    ('Shadow Mode', 'Weeks 3–6', 'Run Lumio side-by-side with current process — no client-facing actions. Compare quality, latency, and failure modes weekly.'),
    ('Limited Production', 'Weeks 7–10', 'Deploy to a controlled user cohort with mandatory human review for high-risk outputs. Enable DTOE learning loop.'),
    ('Scale Decision', 'Weeks 11–12', 'Evaluate pilot KPIs; approve expansion to a second workflow or iterate scope.'),
]

for phase, timeline, activities in pilot_data:
    row = pilot_table.add_row().cells
    row[0].text = phase
    row[1].text = timeline
    row[2].text = activities
    for cell in row:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
    row[0].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

add_rich_para([
    ('Pilot Success Metrics: ', True, False, ACCENT),
    ('Workflow cycle-time reduction · First useful response time · Human rework rate · Citation completeness · Next-best-action adoption uplift vs. non-twin baseline · Policy violation and exception rates.', False, False, GRAY),
], space_after=Pt(12))

# ============================================================
# OUR EDGE
# ============================================================
add_heading_styled('Our Edge', level=1)

add_rich_para([
    ("Lumio AI's differentiation lies in the ", False, False, None),
    ('Digital Twin Optimisation Engine', True, False, None),
    ('. Unlike static rule-based systems, DTOE treats each user or client profile as a probabilistic belief state that evolves with every interaction, decision, and outcome. This creates a compounding advantage: the longer the system operates, the more accurate its recommendations become — delivering measurable, growing value to advisors and their clients.', False, False, None),
])

add_rich_para([
    ('Integration is low-disruption by design. ', True, False, None),
    ('Lumio connects via API to existing CRM, research, ticketing, and approval systems. There is no requirement to replace core infrastructure or existing AI model vendors. The platform coexists with — and enhances — tools already in use.', False, False, None),
])

# Divider
doc.add_paragraph('─' * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# FOOTER
# ============================================================
add_para('Songyi Li', bold=True, size=Pt(11), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para('ucessl7@ucl.ac.uk  ·  Lumio AI  ·  University College London',
         size=Pt(9), color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_para('Prepared for Morgan Stanley',
         italic=True, size=Pt(8.5), color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ---------- save ----------
out_path = os.path.join(os.path.dirname(__file__), 'Lumio_AI_Morgan_Stanley_OnePager.docx')
doc.save(out_path)
print(f'Saved to: {out_path}')
